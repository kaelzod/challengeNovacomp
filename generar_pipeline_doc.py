from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Estilos base ──────────────────────────────────────────────────────────────
def add_title(doc, text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_h2(doc, text):
    doc.add_heading(text, level=2)

def add_h3(doc, text):
    doc.add_heading(text, level=3)

def add_code(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    p.paragraph_format.left_indent = Cm(1)

def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_data in rows:
        row = table.add_row().cells
        for i, text in enumerate(row_data):
            row[i].text = text
            for para in row[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8.5)
    if widths:
        for i, w in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = w
    return table

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
add_title(doc, 'Documentación del Pipeline de Datos')
doc.add_paragraph('Examen Técnico — Data Engineer Semi Senior', style='Subtitle' if 'Subtitle' in [s.name for s in doc.styles] else 'Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
doc.add_paragraph('Fecha: 2026-06-17').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. RESUMEN EJECUTIVO
# ══════════════════════════════════════════════════════════════════════════════
add_h2(doc, '1. Resumen Ejecutivo')
doc.add_paragraph(
    'Este documento describe el pipeline de datos diseñado para el caso técnico de Financiera OH. '
    'El pipeline cubre tres capas: análisis SQL sobre datos de clientes y transacciones, '
    'integración y estandarización de tarjetas de dos procesadores externos, '
    'y validación de calidad sobre un dataset de transacciones anómalas.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. ARQUITECTURA GENERAL
# ══════════════════════════════════════════════════════════════════════════════
add_h2(doc, '2. Arquitectura General del Pipeline')
doc.add_paragraph('El pipeline se divide en 3 partes independientes:')

add_table(doc,
    ['Parte', 'Descripción', 'Entrada', 'Salida'],
    [
        ['Parte 1', 'Análisis SQL de clientes, tarjetas y transacciones', 'customers, cards, transactions', 'Reportes analíticos'],
        ['Parte 2', 'Integración de procesadores y cálculo financiero', 'cards_processor_a/b, equivalence_table_extended, financial_rules', 'cards_financial_master'],
        ['Parte 3', 'Data Quality Report y limpieza de anomalías', 'transactions_anomalies', 'data_quality_report.json\ntransactions_anomalies_clean.csv'],
    ],
    widths=[Cm(2.5), Cm(6), Cm(5), Cm(5)]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 3. FUENTES DE DATOS
# ══════════════════════════════════════════════════════════════════════════════
add_h2(doc, '3. Fuentes de Datos (Tablas de Entrada)')
add_table(doc,
    ['Tabla', 'Columnas Clave', 'Descripción'],
    [
        ['customers', 'customer_id, full_name, segment, register_date', 'Maestro de clientes con segmentación Bronze / Silver / Gold'],
        ['cards', 'card_id, customer_id, card_type, status', 'Tarjetas asociadas a clientes (Mastercard / Visa)'],
        ['transactions', 'tx_id, card_id, amount, tx_date', 'Transacciones financieras por tarjeta'],
        ['cards_processor_a', 'card_id_a, customer_id, card_type_a, status_a', 'Tarjetas del procesador externo A (valores crudos: GOLD_A, ACTIVE_A)'],
        ['cards_processor_b', 'card_id_b, customer_id, card_type_b, status_b', 'Tarjetas del procesador externo B (valores crudos: GOLD_B, ENABLED_B)'],
        ['equivalence_table_extended', 'attribute, value_processor_a, value_processor_b, value_standard', 'Tabla de homologación de valores entre procesadores'],
        ['financial_rules', 'processor, card_type_raw, interest_rate, commission, insurance_fee', 'Reglas financieras por procesador y tipo de tarjeta'],
        ['transactions_anomalies', 'tx_id, card_id, amount, tx_date, extra_column', 'Dataset con anomalías intencionales para validación de calidad'],
        ['business_rules_text', 'rule_id, description', 'Reglas de negocio en lenguaje natural para traducción a SQL'],
    ],
    widths=[Cm(4), Cm(6.5), Cm(8)]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 4. PARTE 1 — ANÁLISIS SQL
# ══════════════════════════════════════════════════════════════════════════════
add_h2(doc, '4. Parte 1 — Análisis SQL')
doc.add_paragraph(
    'Se desarrollaron 6 consultas SQL sobre las tablas customers, cards y transactions. '
    'Las queries utilizan SQL estándar compatible con BigQuery.'
)

add_table(doc,
    ['#', 'Consulta', 'Técnica SQL', 'Resultado'],
    [
        ['1.1', 'Clientes por segmento', 'COUNT + GROUP BY', 'Bronze: 5, Gold: 2, Silver: 3'],
        ['1.2', 'Transacciones y monto por tipo de tarjeta', 'INNER JOIN + SUM + GROUP BY', 'Mastercard $3,853 / Visa $3,549'],
        ['1.3', 'Clientes con tarjeta Blocked', 'INNER JOIN + WHERE status', '3 clientes con tarjeta bloqueada'],
        ['1.4', 'Monto promedio por cliente (>3 tx)', 'Doble JOIN + AVG + HAVING', '3 clientes cumplen el criterio'],
        ['1.5', 'Transacciones sobre percentil 90', 'Subquery + PERCENTILE_CONT', '3 transacciones (>$471)'],
        ['1.6', 'Transacciones con tarjetas inexistentes', 'LEFT JOIN + IS NULL', '0 filas — datos íntegros'],
    ],
    widths=[Cm(1), Cm(5.5), Cm(4.5), Cm(7.5)]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 5. PARTE 2 — PIPELINE cards_financial_master
# ══════════════════════════════════════════════════════════════════════════════
add_h2(doc, '5. Parte 2 — Pipeline cards_financial_master')
doc.add_paragraph(
    'Pipeline de integración construido con 4 CTEs encadenados. '
    'Unifica dos procesadores externos, estandariza valores y calcula costos financieros mensuales.'
)

add_h3(doc, '5.1 Diagrama de flujo')
add_table(doc,
    ['Paso', 'CTE', 'Operación', 'Salida'],
    [
        ['1', 'unified', 'UNION ALL de cards_processor_a y cards_processor_b', 'Estructura común con columna processor (A/B)'],
        ['2', 'standardized', 'Doble LEFT JOIN contra equivalence_table_extended', 'card_type y status en valores estándar (GOLD, ACTIVE…)'],
        ['3', 'enriched', 'LEFT JOIN contra financial_rules por processor + card_type_raw', 'interest_rate, commission, insurance_fee por tarjeta'],
        ['4', 'cards_financial_master', 'Cálculo: interest_rate*1000 + commission + insurance_fee', 'monthly_total_cost + load_date + load_timestamp'],
    ],
    widths=[Cm(1.5), Cm(4), Cm(6), Cm(7)]
)

doc.add_paragraph()
add_h3(doc, '5.2 Decisiones técnicas')
for item in [
    'UNION ALL en lugar de UNION: no existen duplicados entre procesadores, UNION ALL es más eficiente.',
    'LEFT JOIN para estandarización: preserva tarjetas sin equivalencia mapeada para no perder datos.',
    'card_type_raw se conserva en standardized: necesario para el JOIN con financial_rules en el paso siguiente.',
    'load_date y load_timestamp: columnas de auditoría para trazabilidad del pipeline en producción.',
    'En BigQuery productivo: CREATE OR REPLACE TABLE con PARTITION BY load_date y CLUSTER BY processor, customer_id.',
]:
    p = doc.add_paragraph(item, style='List Bullet')
    p.runs[0].font.size = Pt(9)

doc.add_paragraph()
add_h3(doc, '5.3 Muestra del resultado final')
add_table(doc,
    ['card_id', 'customer_id', 'card_type', 'status', 'processor', 'interest_rate', 'commission', 'insurance_fee', 'monthly_total_cost'],
    [
        ['A1001', '1', 'GOLD', 'ACTIVE', 'A', '0.029', '12.5', '5.0', '46.5'],
        ['B2001', '1', 'GOLD', 'ACTIVE', 'B', '0.028', '11.5', '4.5', '44.0'],
        ['A1002', '2', 'SILVER', 'BLOCKED', 'A', '0.035', '10.0', '4.0', '49.0'],
        ['B2003', '3', 'GOLD', 'ACTIVE', 'B', '0.028', '11.5', '4.5', '44.0'],
        ['...', '...', '...', '...', '...', '...', '...', '...', '...'],
    ],
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 6. PARTE 3 — DATA QUALITY
# ══════════════════════════════════════════════════════════════════════════════
add_h2(doc, '6. Parte 3 — Data Quality Report')
doc.add_paragraph(
    'Script Python con pandas que valida transactions_anomalies.csv, '
    'genera un reporte JSON con los problemas detectados y exporta el dataset limpio.'
)

add_h3(doc, '6.1 Validaciones aplicadas')
add_table(doc,
    ['Validación', 'Método', 'Resultado'],
    [
        ['Valores nulos por columna', 'isnull().sum()', 'tx_date: 9 | amount: 5 | card_id: 4 | extra_column: 9'],
        ['Filas duplicadas', 'duplicated().sum()', '0 duplicados'],
        ['Fechas futuras (>hoy)', 'pd.to_datetime + comparación', '5 transacciones con fecha futura'],
        ['card_id inexistentes en procesadores', 'set.union() + isin()', '7 filas inválidas (X9999 + nulls)'],
        ['Campos no documentados', 'set difference de columnas', 'extra_column no está en el esquema'],
        ['Montos negativos', 'amount < 0', '9 transacciones con monto negativo'],
    ],
    widths=[Cm(5), Cm(4.5), Cm(9)]
)

doc.add_paragraph()
add_h3(doc, '6.2 Proceso de limpieza')
for item in [
    'Eliminar filas con nulls en card_id, amount o tx_date.',
    'Eliminar filas con fechas futuras (tx_date > hoy).',
    'Eliminar filas con card_id inexistentes en los procesadores.',
    'Eliminar filas con amount negativo.',
    'Eliminar columna extra_column (no documentada).',
    'Resultado: 1 fila limpia de 20 originales — el dataset era prácticamente todo anomalías.',
]:
    p = doc.add_paragraph(item, style='List Bullet')
    p.runs[0].font.size = Pt(9)

doc.add_paragraph()
add_h3(doc, '6.3 Archivos generados')
add_table(doc,
    ['Archivo', 'Descripción'],
    [
        ['data_quality_report.json', 'Reporte JSON con todos los problemas detectados por categoría'],
        ['transactions_anomalies_clean.csv', 'Dataset limpio listo para carga a BigQuery'],
        ['load_bq.py', 'Script opcional para carga a BigQuery usando google-cloud-bigquery'],
    ],
    widths=[Cm(7), Cm(11.5)]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 7. CONCLUSIONES Y RECOMENDACIONES
# ══════════════════════════════════════════════════════════════════════════════
add_h2(doc, '7. Conclusiones y Recomendaciones')

add_h3(doc, 'Conclusiones')
for item in [
    'El pipeline cubre correctamente las 3 capas: análisis, integración y calidad de datos.',
    'Los datos de la Parte 1 están íntegros: no hay transacciones huérfanas ni inconsistencias graves.',
    'La integración de procesadores requiere tabla de equivalencias para evitar valores crudos en capas de BI.',
    'El dataset de anomalías tiene una tasa de error del 95% — indica que es un archivo de prueba, no productivo.',
]:
    p = doc.add_paragraph(item, style='List Bullet')
    p.runs[0].font.size = Pt(9)

doc.add_paragraph()
add_h3(doc, 'Recomendaciones')
for item in [
    'Implementar la tabla cards_financial_master en BigQuery con PARTITION BY load_date y CLUSTER BY processor, customer_id para optimizar costos.',
    'Formalizar las reglas de negocio (business_rules_text) en tablas estructuradas antes de implementarlas en SQL.',
    'Agregar columna issue_date a cards_processor_b para habilitar la Regla 2 (LEGACY_SILVER).',
    'Definir el valor numérico de "comisión reducida" y "beneficio premium" en financial_rules.',
    'Automatizar el Data Quality Report como paso previo a cualquier carga a BigQuery.',
]:
    p = doc.add_paragraph(item, style='List Bullet')
    p.runs[0].font.size = Pt(9)

doc.save('./archivos/pipeline_documentacion.docx')
print('Documento generado: archivos/pipeline_documentacion.docx')
