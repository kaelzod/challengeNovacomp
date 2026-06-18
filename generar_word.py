from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Título
titulo = doc.add_heading('Desafío Avanzado — Punto B', level=1)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Análisis de inconsistencias en business_rules_text.csv y conversión a reglas SQL.')

doc.add_paragraph()

# Tabla
headers = ['REGLA', 'INTERPRETACIÓN FUNCIONAL', 'TRADUCCIÓN SQL', 'INCONSISTENCIA / LECTURA CORRECTA']

rows = [
    [
        'Regla 1\nGOLD_A migra a GOLD_ESTÁNDAR salvo si BLOCKED_A',
        'Reclasificar tarjetas GOLD_A activas a un producto nuevo.',
        'CASE WHEN card_type_a = \'GOLD_A\'\n  AND status_a != \'BLOCKED_A\'\n  THEN \'GOLD\'\n  ELSE value_standard\nEND AS card_type',
        'GOLD_ESTÁNDAR no existe en equivalence_table_extended. Los valores estándar son GOLD, SILVER, BRONZE. La regla debería decir GOLD.'
    ],
    [
        'Regla 2\nSILVER_B emitida antes de 2020 → LEGACY_SILVER',
        'Identificar tarjetas antiguas y reclasificarlas como producto legacy.',
        '❌ No implementable con los datos actuales.',
        'cards_processor_b no tiene columna de fecha de emisión. Además LEGACY_SILVER no existe en la tabla de equivalencias. Se requiere agregar issue_date.'
    ],
    [
        'Regla 3\nBRONZE con >20 tx mensuales → comisión reducida',
        'Detectar tarjetas de alto uso y aplicar beneficio comercial.',
        'SELECT card_id, COUNT(*) AS tx_mes\nFROM transactions\nGROUP BY card_id,\n  DATE_TRUNC(tx_date, MONTH)\nHAVING COUNT(*) > 20',
        'Dos problemas: (1) cards_processor no tiene historial de transacciones — requiere JOIN con transactions. (2) "comisión reducida" no está cuantificada. Debe definirse en financial_rules.'
    ],
    [
        'Regla 4\nCliente Gold con ≥2 tarjetas activas → beneficio premium',
        'Identificar clientes premium para aplicar ventaja comercial.',
        'SELECT cu.customer_id, cu.full_name\nFROM customers cu\nINNER JOIN cards ca\n  ON cu.customer_id = ca.customer_id\nWHERE cu.segment = \'Gold\'\n  AND ca.status = \'Active\'\nGROUP BY cu.customer_id, cu.full_name\nHAVING COUNT(ca.card_id) >= 2',
        'La detección es implementable, pero "beneficio premium" no está cuantificado. No se sabe si es descuento en tasa, comisión cero, etc. Debe definirse en financial_rules.'
    ],
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Encabezados
hdr = table.rows[0].cells
for i, h in enumerate(headers):
    hdr[i].text = h
    run = hdr[i].paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(9)
    hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Anchos de columna
widths = [Cm(3), Cm(4.5), Cm(5.5), Cm(5.5)]
for i, width in enumerate(widths):
    for cell in table.columns[i].cells:
        cell.width = width

# Filas de datos
for row_data in rows:
    row = table.add_row().cells
    for i, text in enumerate(row_data):
        row[i].text = text
        for para in row[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(8.5)

# Conclusión
doc.add_paragraph()
conclusion = doc.add_heading('Conclusión', level=2)
doc.add_paragraph(
    'De las 4 reglas, solo la Regla 1 y parcialmente la Regla 4 son implementables con los datos disponibles. '
    'Las reglas 2 y 3 tienen dependencias de datos faltantes (fecha de emisión, volumen de transacciones mensual). '
    'Las 4 reglas usan términos no estandarizados (GOLD_ESTÁNDAR, LEGACY_SILVER, comisión reducida, beneficio premium) '
    'que deben formalizarse en las tablas de referencia antes de traducirse a SQL productivo.'
)

doc.save('./archivos/desafio_avanzado_B.docx')
print('Word generado: archivos/desafio_avanzado_B.docx')
